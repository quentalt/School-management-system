import argparse
import csv
import json
import getpass
import os
from datetime import datetime
from docx import Document
from colorama import Fore, Style, init
from typing import List, Dict, Any

# Initialiser colorama
init(autoreset=True)

# Fichiers pour sauvegarder les données
DATA_FILE = "etudiants.json"
DATA_JSON = "data.json"
EXPORT_FILE = "export.txt"

# Couleurs
ROUGE = Fore.RED
VERT = Fore.GREEN
BLEU = Fore.BLUE
JAUNE = Fore.YELLOW
NORMAL = Style.RESET_ALL

# Données en mémoire
etudiants: List[Dict[str, Any]] = []
enseignants: List[Dict[str, Any]] = []
cours: List[Dict[str, Any]] = []
notes: List[Dict[str, Any]] = []
absences: List[Dict[str, Any]] = []

def charger_donnees_fichier(file_path: str) -> Dict[str, Any]:
    """Charge les données depuis un fichier JSON."""
    try:
        with open(file_path, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def sauvegarder_donnees_fichier(file_path: str, data: Dict[str, Any]) -> None:
    """Sauvegarde les données dans un fichier JSON."""
    with open(file_path, "w") as f:
        json.dump(data, f)

def enregistrer_utilisateur() -> None:
    utilisateurs = charger_donnees_fichier(DATA_FILE)
    username = input("Entrez votre nom d'utilisateur : ")
    if username in utilisateurs:
        print("Cet utilisateur existe déjà.")
        return
    password = getpass.getpass("Entrez votre mot de passe : ")
    utilisateurs[username] = password
    sauvegarder_donnees_fichier(DATA_FILE, utilisateurs)
    print("Utilisateur enregistré avec succès!")

def authentifier_utilisateur() -> None:
    utilisateurs = charger_donnees_fichier(DATA_FILE)
    username = input("Entrez votre nom d'utilisateur : ")
    password = getpass.getpass("Entrez votre mot de passe : ")
    if utilisateurs.get(username) == password:
        print("Authentification réussie!")
    else:
        print("Nom d'utilisateur ou mot de passe incorrect.")

def gestion_connexion() -> None:
    while True:
        choix = input("Voulez-vous (1) vous authentifier ou (2) enregistrer un utilisateur? ")
        if choix == "1":
            authentifier_utilisateur()
            break
        elif choix == "2":
            enregistrer_utilisateur()
            break
        else:
            print("Choix invalide. Veuillez choisir une option valide.")

def sauvegarder_etudiants(etudiants: List[Dict[str, Any]]) -> None:
    """Sauvegarde les étudiants dans un fichier JSON."""
    sauvegarder_donnees_fichier(DATA_FILE, etudiants)

def afficher_etudiants(etudiants: List[Dict[str, Any]]) -> None:
    if not etudiants:
        print(Fore.YELLOW + "Aucun étudiant n'est enregistré.")
    else:
        print(Fore.CYAN + "\nListe des étudiants :")
        for idx, etudiant in enumerate(etudiants, 1):
            print(f"{idx}. {etudiant['nom']} - {etudiant['date_ajout']}")
        print()

def rechercher_etudiant_par_date(etudiants: List[Dict[str, Any]]) -> None:
    date_str = input("Entrez la date (format YYYY-MM-DD) à rechercher : ")
    try:
        date_recherche = datetime.strptime(date_str, "%Y-%m-%d")
        resultats = [etudiant for etudiant in etudiants if
                     datetime.strptime(etudiant["date_ajout"], "%Y-%m-%d %H:%M:%S").date() == date_recherche.date()]

        if resultats:
            print(Fore.CYAN + "\nÉtudiants trouvés :")
            for idx, etudiant in enumerate(resultats, 1):
                print(f"{idx}. {etudiant['nom']} - {etudiant['date_ajout']}")
            print()
        else:
            print(Fore.YELLOW + "Aucun étudiant trouvé pour cette date.")
    except ValueError:
        print(Fore.RED + "Format de date invalide. Assurez-vous d'utiliser le format YYYY-MM-DD.")

def exporter_vers_doc(etudiants: List[Dict[str, Any]]) -> None:
    document = Document()
    document.add_heading("Liste des Étudiants", 0)

    if not etudiants:
        document.add_paragraph("Aucun étudiant n'est enregistré.")
    else:
        for etudiant in etudiants:
            document.add_paragraph(f"{etudiant['nom']} - {etudiant['date_ajout']}")

    document.save("etudiants.docx")
    print(Fore.GREEN + "Les étudiants ont été exportés dans le fichier 'etudiants.docx'.")

def exporter_vers_csv(etudiants: List[Dict[str, Any]]) -> None:
    with open("etudiants.csv", "w", newline="") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["Nom", "Date d'ajout"])
        for etudiant in etudiants:
            writer.writerow([etudiant["nom"], etudiant["date_ajout"]])
    print(Fore.GREEN + "Les étudiants ont été exportés dans le fichier 'etudiants.csv'.")

def trier_etudiants(etudiants: List[Dict[str, Any]]) -> None:
    critere = input("Trier par (1) nom (2) date d'ajout : ")
    if critere == "1":
        etudiants.sort(key=lambda e: e["nom"].lower())
        print(Fore.GREEN + "Les étudiants ont été triés par nom.")
    elif critere == "2":
        etudiants.sort(key=lambda e: datetime.strptime(e["date_ajout"], "%Y-%m-%d %H:%M:%S"))
        print(Fore.GREEN + "Les étudiants ont été triés par date d'ajout.")
    else:
        print(Fore.YELLOW + "Critère invalide.")

def sauvegarder_donnees() -> None:
    data = {
        "etudiants": etudiants,
        "enseignants": enseignants,
        "cours": cours,
        "notes": notes,
        "absences": absences
    }
    sauvegarder_donnees_fichier(DATA_JSON, data)
    print(VERT + "Données sauvegardées avec succès!" + NORMAL)

def charger_donnees() -> None:
    global etudiants, enseignants, cours, notes, absences
    data = charger_donnees_fichier(DATA_JSON)
    etudiants = data.get("etudiants", [])
    enseignants = data.get("enseignants", [])
    cours = data.get("cours", [])
    notes = data.get("notes", [])
    absences = data.get("absences", [])
    print(VERT + "Données chargées avec succès!" + NORMAL)

def afficher_statistiques() -> None:
    print(BLEU + "Statistiques:" + NORMAL)
    print(f"{VERT}Nombre d'étudiants : {len(etudiants)}")
    print(f"Nombre d'enseignants : {len(enseignants)}")
    print(f"Nombre de cours : {len(cours)}")
    print(f"Nombre de notes : {len(notes)}")
    print(f"Nombre d'absences : {len(absences)}" + NORMAL)

def exporter_donnees() -> None:
    with open(EXPORT_FILE, "w") as f:
        f.write("Étudiants:\n")
        for etudiant in etudiants:
            f.write(f"{etudiant}\n")
        f.write("\nEnseignants:\n")
        for enseignant in enseignants:
            f.write(f"{enseignant}\n")
        f.write("\nCours:\n")
        for c in cours:
            f.write(f"{c}\n")
        f.write("\nNotes:\n")
        for note in notes:
            f.write(f"{note}\n")
        f.write("\nAbsences:\n")
        for absence in absences:
            f.write(f"{absence}\n")
    print(VERT + "Données exportées dans 'export.txt' avec succès!" + NORMAL)


def rechercher_etudiant_par_nom(etudiants: List[Dict[str, Any]]) -> None:
    nom = input("Entrez le nom de l'étudiant à rechercher : ")
    resultats = [etudiant for etudiant in etudiants if etudiant["nom"].lower() == nom.lower()]
    if resultats:
        print(Fore.CYAN + "\nÉtudiants trouvés :")
        for idx, etudiant in enumerate(resultats, 1):
            date_ajout = etudiant.get("date_ajout", "Date non disponible")
            print(f"{idx}. {etudiant['nom']} - {date_ajout}")
        print()
    else:
        print(Fore.YELLOW + "Aucun étudiant trouvé pour ce nom.")

def rechercher_enseignant_par_nom(enseignants: List[Dict[str, Any]]) -> None:
    nom = input("Entrez le nom de l'enseignant à rechercher : ")
    resultats = [enseignant for enseignant in enseignants if enseignant["nom"].lower() == nom.lower()]
    if resultats:
        print(Fore.CYAN + "\nEnseignants trouvés :")
        for idx, enseignant in enumerate(resultats, 1):
            print(f"{idx}. {enseignant['nom']}")
        print()
    else:
        print(Fore.YELLOW + "Aucun enseignant trouvé pour ce nom.")

def rechercher_cours_par_nom(cours: List[Dict[str, Any]]) -> None:
    nom = input("Entrez le nom du cours à rechercher : ")
    resultats = [c for c in cours if c["nom"].lower() == nom.lower()]
    if resultats:
        print(Fore.CYAN + "\nCours trouvés :")
        for idx, c in enumerate(resultats, 1):
            print(f"{idx}. {c['nom']}")
        print()
    else:
        print(Fore.YELLOW + "Aucun cours trouvé pour ce nom.")

def rechercher_note_par_etudiant(notes: List[Dict[str, Any]]) -> None:
    etudiant = input("Entrez le nom de l'étudiant : ")
    resultats = [note for note in notes if note["etudiant"].lower() == etudiant.lower()]
    if resultats:
        print(Fore.CYAN + "\nNotes trouvées :")
        for idx, note in enumerate(resultats, 1):
            print(f"{idx}. {note['etudiant']} - {note['note']}")
        print()
    else:
        print(Fore.YELLOW + "Aucune note trouvée pour cet étudiant.")

def rechercher() -> None:
    print(BLEU + "Rechercher:" + NORMAL)
    print(VERT + "1. Rechercher un étudiant par date d'ajout" + NORMAL)
    print(VERT + "2. Rechercher un étudiant par nom" + NORMAL)
    print(JAUNE + "3. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        rechercher_etudiant_par_date(etudiants)
    elif choix == "2":
        rechercher_etudiant_par_nom(etudiants)
    elif choix == "3":
        return
    else:
        print(ROUGE + "Option invalide. Veuillez choisir une option valide." + NORMAL)


def afficher_menu() -> None:
    print(ROUGE + "Menu:" + NORMAL)
    print(VERT + "1. Gestion des étudiants" + NORMAL)
    print(VERT + "2. Gestion des enseignants" + NORMAL)
    print(VERT + "3. Gestion des cours" + NORMAL)
    print(VERT + "4. Gestion des notes" + NORMAL)
    print(VERT + "5. Gestion des absences" + NORMAL)
    print(VERT + "6. Lister les étudiants" + NORMAL)
    print(VERT + "7. Lister les enseignants" + NORMAL)
    print(VERT + "8. Lister les absences" + NORMAL)
    print(BLEU + "9. Rechercher" + NORMAL)
    print(BLEU + "10. Statistiques" + NORMAL)
    print(BLEU + "11. Exporter les données" + NORMAL)
    print(JAUNE + "12. Sauvegarder" + NORMAL)
    print(JAUNE + "13. Charger" + NORMAL)
    print(JAUNE + "14. Quitter" + NORMAL)

def gestion_etudiants() -> None:
    print(ROUGE + "Gestion des étudiants:" + NORMAL)
    print(VERT + "1. Ajouter un étudiant" + NORMAL)
    print(VERT + "2. Modifier un étudiant" + NORMAL)
    print(VERT + "3. Supprimer un étudiant" + NORMAL)
    print(JAUNE + "4. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        ajouter_etudiant()
    elif choix == "2":
        modifier_etudiant()
    elif choix == "3":
        supprimer_etudiant()
    elif choix == "4":
        return
    else:
        print("Option invalidée. Veuillez choisir une option valide.")

def gestion_enseignants() -> None:
    print(ROUGE + "Gestion des enseignants:" + NORMAL)
    print(VERT + "1. Ajouter un enseignant" + NORMAL)
    print(VERT + "2. Modifier un enseignant" + NORMAL)
    print(VERT + "3. Supprimer un enseignant" + NORMAL)
    print(JAUNE + "4. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        ajouter_enseignant()
    elif choix == "2":
        modifier_enseignant()
    elif choix == "3":
        supprimer_enseignant()
    elif choix == "4":
        return
    else:
        print("Option invalidée. Veuillez choisir une option valide.")

def gestion_cours() -> None:
    print(ROUGE + "Gestion des cours:" + NORMAL)
    print(VERT + "1. Ajouter un cours" + NORMAL)
    print(VERT + "2. Modifier un cours" + NORMAL)
    print(VERT + "3. Supprimer un cours" + NORMAL)
    print(JAUNE + "4. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        ajouter_cours()
    elif choix == "2":
        modifier_cours()
    elif choix == "3":
        supprimer_cours()
    elif choix == "4":
        return
    else:
        print("Option invalidée. Veuillez choisir une option valide.")

def gestion_notes() -> None:
    print(ROUGE + "Gestion des notes:" + NORMAL)
    print(VERT + "1. Ajouter une note" + NORMAL)
    print(VERT + "2. Modifier une note" + NORMAL)
    print(VERT + "3. Supprimer une note" + NORMAL)
    print(JAUNE + "4. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        ajouter_note()
    elif choix == "2":
        modifier_note()
    elif choix == "3":
        supprimer_note()
    elif choix == "4":
        return
    else:
        print("Option invalidée. Veuillez choisir une option valide.")

def gestion_absences() -> None:
    print(ROUGE + "Gestion des absences:" + NORMAL)
    print(VERT + "1. Ajouter une absence" + NORMAL)
    print(VERT + "2. Modifier une absence" + NORMAL)
    print(VERT + "3. Supprimer une absence" + NORMAL)
    print(JAUNE + "4. Retour" + NORMAL)
    choix = input("Choisissez une option : ")
    if choix == "1":
        ajouter_absence()
    elif choix == "2":
        modifier_absence()
    elif choix == "3":
        supprimer_absence()
    elif choix == "4":
        return
    else:
        print("Option invalidée. Veuillez choisir une option valide.")

def lister_etudiants() -> None:
    print(BLEU + "Liste des étudiants:" + NORMAL)
    for etudiant in etudiants:
        print(etudiant)

def lister_enseignants() -> None:
    print(BLEU + "Liste des enseignants:" + NORMAL)
    for enseignant in enseignants:
        print(enseignant)

def lister_absences() -> None:
    print(BLEU + "Liste des absences:" + NORMAL)
    for absence in absences:
        print(absence)

def ajouter_etudiant() -> None:
    nom = input("Entrez le nom de l'étudiant : ")
    if any(etudiant["nom"].lower() == nom.lower() for etudiant in etudiants):
        print(Fore.YELLOW + "Cet étudiant existe déjà.")
    else:
        date_ajout = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        etudiants.append({"nom": nom, "date_ajout": date_ajout})
        print(Fore.GREEN + f"{nom} a été ajouté avec succès.")

def modifier_etudiant() -> None:
    nom = input("Entrez le nom de l'étudiant à modifier : ")
    for etudiant in etudiants:
        if etudiant["nom"] == nom:
            nouveau_nom = input("Entrez le nouveau nom de l'étudiant : ")
            etudiant["nom"] = nouveau_nom
            print(VERT + "Étudiant modifié avec succès!" + NORMAL)
            return
    print(ROUGE + "Étudiant non trouvé." + NORMAL)

def supprimer_etudiant() -> None:
    nom = input("Entrez le nom de l'étudiant à supprimer : ")
    for etudiant in etudiants:
        if etudiant["nom"] == nom:
            etudiants.remove(etudiant)
            print(VERT + "Étudiant supprimé avec succès!" + NORMAL)
            return
    print(ROUGE + "Étudiant non trouvé." + NORMAL)

def ajouter_enseignant() -> None:
    nom = input("Entrez le nom de l'enseignant : ")
    enseignants.append({"nom": nom})
    print(VERT + "Enseignant ajouté avec succès!" + NORMAL)

def modifier_enseignant() -> None:
    nom = input("Entrez le nom de l'enseignant à modifier : ")
    for enseignant in enseignants:
        if enseignant["nom"] == nom:
            nouveau_nom = input("Entrez le nouveau nom de l'enseignant : ")
            enseignant["nom"] = nouveau_nom
            print(VERT + "Enseignant modifié avec succès!" + NORMAL)
            return
    print(ROUGE + "Enseignant non trouvé." + NORMAL)

def supprimer_enseignant() -> None:
    nom = input("Entrez le nom de l'enseignant à supprimer : ")
    for enseignant in enseignants:
        if enseignant["nom"] == nom:
            enseignants.remove(enseignant)
            print(VERT + "Enseignant supprimé avec succès!" + NORMAL)
            return
    print(ROUGE + "Enseignant non trouvé." + NORMAL)

def ajouter_cours() -> None:
    nom = input("Entrez le nom du cours : ")
    cours.append({"nom": nom})
    print(VERT + "Cours ajouté avec succès!" + NORMAL)

def modifier_cours() -> None:
    nom = input("Entrez le nom du cours à modifier : ")
    for c in cours:
        if c["nom"] == nom:
            nouveau_nom = input("Entrez le nouveau nom du cours : ")
            c["nom"] = nouveau_nom
            print(VERT + "Cours modifié avec succès!" + NORMAL)
            return
    print(ROUGE + "Cours non trouvé." + NORMAL)

def supprimer_cours() -> None:
    nom = input("Entrez le nom du cours à supprimer : ")
    for c in cours:
        if c["nom"] == nom:
            cours.remove(c)
            print(VERT + "Cours supprimé avec succès!" + NORMAL)
            return
    print(ROUGE + "Cours non trouvé." + NORMAL)

def ajouter_note() -> None:
    etudiant_nom = input("Entrez le nom de l'étudiant : ")
    note = input("Entrez la note : ")
    notes.append({"etudiant": etudiant_nom, "note": note})
    print(VERT + "Note ajoutée avec succès!" + NORMAL)

def modifier_note() -> None:
    etudiant_nom = input("Entrez le nom de l'étudiant : ")
    for note in notes:
        if note["etudiant"] == etudiant_nom:
            nouvelle_note = input("Entrez la nouvelle note : ")
            note["note"] = nouvelle_note
            print(VERT + "Note modifiée avec succès!" + NORMAL)
            return
    print(ROUGE + "Note non trouvée." + NORMAL)
def supprimer_note() -> None:
            etudiant_nom = input("Entrez le nom de l'étudiant : ")
            for note in notes:
                if note["etudiant"] == etudiant_nom:
                    notes.remove(note)
                    print(VERT + "Note supprimée avec succès!" + NORMAL)
                    return
            print(ROUGE + "Note non trouvée." + NORMAL)

def ajouter_absence() -> None:
            etudiant_nom = input("Entrez le nom de l'étudiant : ")
            date_absence = input("Entrez la date de l'absence (format YYYY-MM-DD) : ")
            absences.append({"etudiant": etudiant_nom, "date": date_absence})
            print(VERT + "Absence ajoutée avec succès!" + NORMAL)

def modifier_absence() -> None:
            etudiant_nom = input("Entrez le nom de l'étudiant : ")
            for absence in absences:
                if absence["etudiant"] == etudiant_nom:
                    nouvelle_date = input("Entrez la nouvelle date de l'absence (format YYYY-MM-DD) : ")
                    absence["date"] = nouvelle_date
                    print(VERT + "Absence modifiée avec succès!" + NORMAL)
                    return
            print(ROUGE + "Absence non trouvée." + NORMAL)
def supprimer_absence() -> None:
            etudiant_nom = input("Entrez le nom de l'étudiant : ")
            for absence in absences:
                if absence["etudiant"] == etudiant_nom:
                    absences.remove(absence)
                    print(VERT + "Absence supprimée avec succès!" + NORMAL)
                    return
            print(ROUGE + "Absence non trouvée." + NORMAL)

def main() -> None:
            charger_donnees()
            gestion_connexion()
            while True:
                afficher_menu()
                choix = input("Choisissez une option : ")
                if choix == "1":
                    gestion_etudiants()
                elif choix == "2":
                    gestion_enseignants()
                elif choix == "3":
                    gestion_cours()
                elif choix == "4":
                    gestion_notes()
                elif choix == "5":
                    gestion_absences()
                elif choix == "6":
                    lister_etudiants()
                elif choix == "7":
                    lister_enseignants()
                elif choix == "8":
                    lister_absences()
                elif choix == "9":
                    rechercher()
                elif choix == "10":
                    afficher_statistiques()
                elif choix == "11":
                    exporter_donnees()
                elif choix == "12":
                    sauvegarder_donnees()
                elif choix == "13":
                    charger_donnees()
                elif choix == "14":
                    print(JAUNE + "Au revoir!" + NORMAL)
                    break
                else:
                    print(ROUGE + "Option invalide. Veuillez choisir une option valide." + NORMAL)

if __name__ == "__main__":
            parser = argparse.ArgumentParser(description="Gestionnaire d'étudiants")
            parser.add_argument("--export", action="store_true", help="Exporter les données")
            args = parser.parse_args()
            if args.export:
                exporter_donnees()
            else:
                main()